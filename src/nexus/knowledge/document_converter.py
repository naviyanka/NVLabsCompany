"""Document conversion parsers for PDF and DOCX files.

Extends the RAG pipeline's DocumentParser protocol with converters that
extract text from binary document formats and produce ParsedChunks ready
for chunking/embedding.

Requires: pypdfium2, python-docx (optional — graceful import failure).
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from typing import Any

from nexus.knowledge.parsers import DocumentParser, ParsedChunk

logger = logging.getLogger(__name__)


@dataclass
class PDFParser:
    """Extracts text from PDF files page-by-page using pypdfium2.

    Each page becomes one ParsedChunk with page_number metadata.
    Large pages are further split at paragraph boundaries.
    """

    max_chunk_size: int = 1500

    def parse(self, content: bytes | str) -> list[ParsedChunk]:
        import pypdfium2 as pdfium

        if isinstance(content, str):
            content = content.encode("latin-1")

        pdf = pdfium.PdfDocument(content)
        chunks: list[ParsedChunk] = []

        for page_idx in range(len(pdf)):
            page = pdf[page_idx]
            text = page.get_textpage().get_text_range()
            if not text or not text.strip():
                continue

            if len(text) <= self.max_chunk_size:
                chunks.append(ParsedChunk(
                    content=text.strip(),
                    metadata={"page_number": page_idx + 1, "source_type": "pdf"},
                    chunk_type="page",
                ))
            else:
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                for para_idx, para in enumerate(paragraphs):
                    chunks.append(ParsedChunk(
                        content=para,
                        metadata={
                            "page_number": page_idx + 1,
                            "paragraph_index": para_idx,
                            "source_type": "pdf",
                        },
                        chunk_type="paragraph",
                    ))

        return chunks

    def parse_file(self, file_path: str) -> list[ParsedChunk]:
        with open(file_path, "rb") as f:
            return self.parse(f.read())


@dataclass
class DOCXParser:
    """Extracts text from DOCX files paragraph-by-paragraph using python-docx.

    Each paragraph becomes one ParsedChunk. Headings get chunk_type="heading"
    with heading level in metadata.
    """

    max_chunk_size: int = 1500

    def parse(self, content: bytes | str) -> list[ParsedChunk]:
        import io
        from docx import Document

        if isinstance(content, str):
            content = content.encode("latin-1")

        doc = Document(io.BytesIO(content))
        chunks: list[ParsedChunk] = []
        current_text = ""
        current_meta: dict[str, Any] = {"source_type": "docx"}

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                if current_text:
                    chunks.append(ParsedChunk(
                        content=current_text.strip(),
                        metadata=dict(current_meta, paragraph_index=para_idx),
                        chunk_type="paragraph",
                    ))
                    current_text = ""
                continue

            if para.style and para.style.name.startswith("Heading"):
                if current_text:
                    chunks.append(ParsedChunk(
                        content=current_text.strip(),
                        metadata=dict(current_meta, paragraph_index=para_idx),
                        chunk_type="paragraph",
                    ))
                    current_text = ""

                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                chunks.append(ParsedChunk(
                    content=text,
                    metadata={"heading_level": int(level) if level.isdigit() else 1, "source_type": "docx"},
                    chunk_type="heading",
                ))
            elif len(current_text) + len(text) > self.max_chunk_size:
                if current_text:
                    chunks.append(ParsedChunk(
                        content=current_text.strip(),
                        metadata=dict(current_meta, paragraph_index=para_idx),
                        chunk_type="paragraph",
                    ))
                current_text = text + "\n"
            else:
                current_text += text + "\n"

        if current_text.strip():
            chunks.append(ParsedChunk(
                content=current_text.strip(),
                metadata=dict(current_meta),
                chunk_type="paragraph",
            ))

        return chunks

    def parse_file(self, file_path: str) -> list[ParsedChunk]:
        with open(file_path, "rb") as f:
            return self.parse(f.read())


def get_parser_for_extension(ext: str) -> PDFParser | DOCXParser | None:
    """Return appropriate parser for file extension, or None if unsupported."""
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        return PDFParser()
    elif ext in ("docx", "doc"):
        return DOCXParser()
    return None
